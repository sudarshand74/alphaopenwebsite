from pathlib import Path
from datetime import date
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "AlphaOpen_Architecture_Review_Board_Detailed_Design.docx"
ARTIFACTS = ROOT / "output" / "arb_design"
ARTIFACTS.mkdir(parents=True, exist_ok=True)

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF5D6"
PALE_RED = "FDECEC"
PALE_GREEN = "EAF5EC"
WHITE = "FFFFFF"
BLACK = "111111"
RISK_RED = "9B1C1C"
CAUTION = "7A5A00"
GREEN = "256D3B"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run(run, size=11, bold=False, italic=False, color=BLACK, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="CDD3DB", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths, header_fill=LIGHT, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        set_run(p.add_run(str(value)), size=font_size, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run(p.add_run(str(value)), size=font_size, color=BLACK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=NAVY)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_run(p.add_run(text))
    return p


def begin_numbering(doc):
    numbering = doc.part.numbering_part.element
    style_num_id = str(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    abstract_id = "0"
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) == style_num_id:
            abstract = num.find(qn("w:abstractNumId"))
            if abstract is not None:
                abstract_id = abstract.get(qn("w:val"))
            break
    existing_ids = [int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))]
    new_id = max(existing_ids or [0]) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    new_num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    new_num.append(level_override)
    numbering.append(new_num)
    return new_id


def add_number(doc, text, num_id):
    p = doc.add_paragraph(style="List Number")
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id
    set_run(p.add_run(text))
    return p


def add_picture_with_alt(doc, path, alt_text, width):
    shape = doc.add_picture(str(path), width=width)
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)
    return shape


def add_callout(doc, label, text, fill=PALE_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"{label}: "), bold=True, color=color)
    set_run(p.add_run(text), color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_page_break(doc):
    # Heading styles use keep-with-next, allowing Word to paginate naturally
    # without leaving nearly empty pages when the prior table flows.
    return None


def add_diagram(title, boxes, arrows, filename):
    width, height = 1600, 720
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 38)
        box_font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 25)
        small_font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 20)
    except OSError:
        title_font = box_font = small_font = ImageFont.load_default()
    draw.text((50, 28), title, fill=f"#{NAVY}", font=title_font)
    for box in boxes:
        x1, y1, x2, y2, label, fill = box
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=f"#{fill}", outline=f"#{BLUE}", width=3)
        lines = label.split("\n")
        total_h = len(lines) * 34
        y = (y1 + y2 - total_h) / 2
        for line in lines:
            bounds = draw.textbbox((0, 0), line, font=box_font)
            text_w = bounds[2] - bounds[0]
            draw.text(((x1 + x2 - text_w) / 2, y), line, fill=f"#{NAVY}", font=box_font)
            y += 34
    for x1, y1, x2, y2, label in arrows:
        draw.line((x1, y1, x2, y2), fill=f"#{DARK_BLUE}", width=5)
        angle = 12
        if abs(x2 - x1) >= abs(y2 - y1):
            direction = 1 if x2 > x1 else -1
            draw.polygon([(x2, y2), (x2 - direction * 18, y2 - 11), (x2 - direction * 18, y2 + 11)], fill=f"#{DARK_BLUE}")
        else:
            direction = 1 if y2 > y1 else -1
            draw.polygon([(x2, y2), (x2 - 11, y2 - direction * 18), (x2 + 11, y2 - direction * 18)], fill=f"#{DARK_BLUE}")
        if label:
            draw.text(((x1 + x2) / 2 - 55, (y1 + y2) / 2 - 28), label, fill=f"#{MUTED}", font=small_font)
    path = ARTIFACTS / filename
    image.save(path)
    return path


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    list2 = doc.styles["List Bullet 2"]
    list2.font.name = "Calibri"
    list2.font.size = Pt(11)
    list2.paragraph_format.left_indent = Inches(0.75)
    list2.paragraph_format.first_line_indent = Inches(-0.25)
    list2.paragraph_format.space_after = Pt(6)
    list2.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("ALPHAOPEN | ARCHITECTURE REVIEW BOARD"), size=8.5, bold=True, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("Architecture Detailed Design | July 2026 | "), size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def build():
    doc = Document()
    configure_document(doc)

    # Opening block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("ARCHITECTURE DETAILED DESIGN"), size=23, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    set_run(p.add_run("AlphaOpen Tennis League Progressive Web Application"), size=15, color=MUTED)
    metadata = [
        ("Document status", "Proposed for Architecture Review Board"),
        ("Version", "1.0"),
        ("Review date", "July 26, 2026"),
        ("System", "AlphaOpen Tennis League"),
        ("Production project", "alphaopen-development-2026"),
        ("Current hosting", "https://alphaopen-development-2026.web.app"),
        ("Architecture posture", "Firebase Spark-safe, client-mediated transactions"),
    ]
    add_table(doc, ["Control", "Value"], metadata, [2700, 6660], header_fill=PALE_BLUE, font_size=9.5)
    add_callout(
        doc,
        "ARB recommendation",
        "Approve the current Spark-safe architecture as an interim operating model, subject to the controls and remediation plan in Section 15. The design is suitable for a trusted, small league administration model but is not the recommended final state for high-assurance or materially scaled operations.",
        fill=PALE_GOLD,
        color=CAUTION,
    )

    add_heading(doc, "Decision requested", 1)
    for item in [
        "Approve the current Firebase Hosting, Authentication, Firestore, and PWA architecture for controlled production use.",
        "Approve the team-level lineup workflow and the Spark-safe transaction pattern as the current system of record behavior.",
        "Accept the documented residual risks of privileged client execution, especially exceptional lineup reset validation.",
        "Approve the future Blaze migration pattern without requiring a Firestore data migration.",
        "Assign owners and target dates for the conditions of approval.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Document purpose and evidence base", 1)
    add_body(
        doc,
        "This document describes the as-built application architecture, data model, security model, operational workflows, deployment topology, and planned evolution. It is based on the deployed source configuration and the repository artifacts ALPHAOPEN_APP_SPEC.md, FIRESTORE_DATA_MODEL.md, firebase.json, firestore.rules, firestore.indexes.json, manifest.webmanifest, the route and feature modules, and the Spark-safe lineup workflow implementation.",
    )

    add_page_break(doc)
    add_heading(doc, "1. Executive architecture summary", 1)
    add_body(
        doc,
        "AlphaOpen is a mobile-first, single-page Progressive Web Application for public league information and authenticated league operations. Static HTML, CSS, JavaScript modules, images, and the service worker are served by Firebase Hosting. Google Authentication establishes identity. Cloud Firestore is the operational and public data store. Firestore Security Rules are the primary authorization and workflow enforcement layer.",
    )
    add_body(
        doc,
        "The application currently operates on the Firebase Spark plan. There is no deployed trusted application server. Standard create, update, approval, publication, and reset actions execute in the browser using Firestore transactions or batched writes. The security design therefore combines role data, immutable snapshots on matchup records, strict document transition rules, and create-only audit/revision records.",
    )
    summary_rows = [
        ("Presentation", "Responsive SPA/PWA using HTML, CSS, JavaScript modules and route-based lazy loading", "Production"),
        ("Identity", "Google Sign-In through Firebase Authentication; verified-email player linking and admin approval", "Production"),
        ("Application services", "Browser modules and Firestore transactions; no deployed Cloud Functions", "Interim"),
        ("Data", "Cloud Firestore, default database, nam5 multi-region", "Production"),
        ("Authorization", "Firestore Security Rules plus client route visibility", "Production"),
        ("Hosting", "Firebase Hosting with SPA rewrite and no-cache code headers", "Production"),
        ("Offline", "Service worker app-shell cache; operational writes still require Firestore connectivity", "Limited"),
        ("Delivery", "Manual Firebase CLI deployment to one Firebase project", "Needs hardening"),
    ]
    add_table(doc, ["Layer", "Current design", "Assessment"], summary_rows, [1500, 5950, 1910], font_size=8.7)

    add_heading(doc, "Architecture quality assessment", 2)
    quality_rows = [
        ("Strengths", "Simple deployment, low operating cost, responsive PWA, direct real-time reads, strong document ownership model, immutable lineup revisions, transaction-based concurrency."),
        ("Constraints", "No trusted execution tier, no scheduled/background processing, no centralized secrets, limited server-side validation, manual cache versioning, one environment."),
        ("Highest residual risk", "A privileged user can attempt direct SDK writes. Rules control status transitions, but application-level checks such as score inspection during exceptional reset cannot be independently queried by Firestore Rules."),
        ("Recommended disposition", "Approve with conditions and define explicit triggers for a Blaze/server migration."),
    ]
    add_table(doc, ["Category", "Assessment"], quality_rows, [2000, 7360], header_fill=LIGHT, font_size=9)

    add_page_break(doc)
    add_heading(doc, "2. Scope, principles, and design decisions", 1)
    add_heading(doc, "2.1 In-scope capabilities", 2)
    for item in [
        "Public home, season, match, standings, history, rules, venue, and community content.",
        "Google authentication, user registration, Player Master identity matching, and Super Admin approval.",
        "Season, team, roster, venue, matchup, approver, and player administration.",
        "Captain, EC, Neutral Approver, and Super Admin lineup submission.",
        "Independent Home/Away team-level approval, rejection, and self-approval tracking.",
        "Full lineup publication, match scheduling, score entry, calculated points, and standings publication.",
        "Exceptional reset of both fully approved lineups when no score activity is detected.",
        "PWA installation and app-shell offline behavior.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 Governing principles", 2)
    principles = [
        ("Canonical identity", "Firebase UID identifies the account; Player ID identifies the person across seasons and history."),
        ("Season isolation", "Operational data is nested below a season and is governed by active season membership."),
        ("Least privilege", "Guest, Player, Captain, EC, Neutral Approver, and Super Admin access is evaluated separately."),
        ("Immutable history", "Submitted lineup revisions and workflow actions are create-only."),
        ("Atomic state change", "Cross-document workflow state is committed in one transaction or batch."),
        ("Public/private separation", "Public projections are separate from operational and private player records."),
        ("Snapshot history", "Names, ranks, teams, and actor details are copied into official records to preserve historical meaning."),
        ("Plan portability", "UI and schema contracts remain stable when protected mutations move to Cloud Functions."),
    ]
    add_table(doc, ["Principle", "Application"], principles, [2100, 7260], font_size=9)

    add_heading(doc, "2.3 Key architecture decisions", 2)
    decisions = [
        ("ADR-001", "Progressive Web App", "One responsive codebase for phone, tablet, and desktop; avoids native app distribution."),
        ("ADR-002", "Firebase managed platform", "Use Hosting, Authentication, and Firestore to minimize operational overhead."),
        ("ADR-003", "Spark-safe workflow", "Use browser-side Firestore transactions while Blaze is unavailable."),
        ("ADR-004", "Team-level lineup states", "Track Home and Away independently and derive matchup approval status."),
        ("ADR-005", "Immutable lineup revisions", "Never overwrite submitted history; current lineup points to a numbered revision."),
        ("ADR-006", "Separate reset exception", "Fully approved lineups cannot be changed through normal approval; reset always affects both."),
        ("ADR-007", "On-demand score inspection", "Do not maintain scoreActivityStarted; query and re-read line records only when reset is requested."),
        ("ADR-008", "Future trusted service seam", "Expose workflow operations through one client service so Cloud Functions can replace transaction internals."),
    ]
    add_table(doc, ["ID", "Decision", "Rationale"], decisions, [1100, 2600, 5660], font_size=8.6)

    add_page_break(doc)
    add_heading(doc, "3. System context", 1)
    context = add_diagram(
        "AlphaOpen system context",
        [
            (50, 170, 330, 310, "Guests", PALE_BLUE),
            (50, 390, 330, 530, "Players / Captains\nEC / Approvers", PALE_GREEN),
            (570, 220, 1030, 480, "AlphaOpen PWA\nFirebase Hosting", LIGHT),
            (1270, 95, 1540, 235, "Firebase\nAuthentication", PALE_BLUE),
            (1270, 290, 1540, 430, "Cloud\nFirestore", PALE_BLUE),
            (860, 535, 1160, 675, "Google Identity\nProvider", PALE_BLUE),
        ],
        [
            (330, 240, 570, 300, "HTTPS"),
            (330, 460, 570, 400, "HTTPS"),
            (1030, 275, 1270, 165, "Auth SDK"),
            (1030, 350, 1270, 360, "Firestore SDK"),
            (1160, 595, 1325, 235, "OIDC"),
        ],
        "system_context.png",
    )
    add_picture_with_alt(
        doc,
        context,
        "AlphaOpen system context showing guests, authenticated users, Firebase Hosting, Authentication, Firestore, and Google identity provider.",
        Inches(6.5),
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(caption.add_run("Figure 1. System context and managed-service boundaries"), size=9, italic=True, color=MUTED)

    add_heading(doc, "3.1 Actors and trust boundaries", 2)
    actor_rows = [
        ("Guest", "Untrusted anonymous browser", "Published/public documents only"),
        ("Player", "Authenticated active user", "Own account, membership, availability, permitted history"),
        ("Captain", "Authenticated operational user", "Own team lineup, scheduling, and score responsibilities"),
        ("EC", "Season administrator", "Rosters, season structure, venues, disputes, operational visibility"),
        ("Neutral Approver", "Privileged seasonal assignment", "Either team submission; independent approval/rejection; reset exception"),
        ("Super Admin", "Highest privilege", "Global administration and all operational actions"),
        ("Firebase services", "Managed trusted platform", "Identity tokens, rule evaluation, transactions, persistence, hosting"),
    ]
    add_table(doc, ["Actor", "Trust classification", "Primary access"], actor_rows, [1650, 2500, 5210], font_size=8.8)

    add_heading(doc, "3.2 External dependencies", 2)
    for item in [
        "Google Identity Provider for interactive account authentication.",
        "Firebase JavaScript SDK version 12.16.0 loaded from the Google CDN.",
        "Firebase Hosting, Authentication, and Cloud Firestore in project alphaopen-development-2026.",
        "Browser platform services: service workers, Cache API, local/session persistence, dialogs, canvas, and clipboard.",
        "Firebase CLI for rules and hosting deployment.",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "4. Logical application architecture", 1)
    logical = add_diagram(
        "Logical component architecture",
        [
            (40, 150, 360, 275, "Presentation\nindex.html + styles", LIGHT),
            (40, 375, 360, 500, "PWA Runtime\nservice worker + manifest", LIGHT),
            (560, 95, 1010, 220, "Routing and Shell\napp.js + runtime-loader", PALE_BLUE),
            (560, 280, 1010, 405, "Domain Feature Modules\nlineup, roster, scores, admin", PALE_GREEN),
            (560, 465, 1010, 590, "Workflow Services\ntransactions + validation", PALE_GOLD),
            (1220, 150, 1550, 275, "Firebase Auth", PALE_BLUE),
            (1220, 375, 1550, 500, "Cloud Firestore\nRules + Indexes", PALE_BLUE),
        ],
        [
            (360, 210, 560, 160, ""),
            (360, 438, 560, 345, ""),
            (785, 220, 785, 280, ""),
            (785, 405, 785, 465, ""),
            (1010, 160, 1220, 210, ""),
            (1010, 345, 1220, 438, ""),
            (1010, 530, 1220, 438, ""),
        ],
        "logical_architecture.png",
    )
    add_picture_with_alt(
        doc,
        logical,
        "Logical component architecture showing presentation, routing, domain modules, workflow services, authentication, Firestore, and PWA runtime.",
        Inches(6.5),
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(caption.add_run("Figure 2. Browser modules and Firebase service boundaries"), size=9, italic=True, color=MUTED)

    add_heading(doc, "4.1 Component catalog", 2)
    components = [
        ("Application shell", "index.html, styles.css, app.js", "Routes, navigation, responsive views, dialogs, common rendering."),
        ("Runtime loader", "runtime-loader.js and feature bootstraps", "Loads core Firebase modules and route-specific modules on demand."),
        ("Identity", "firebase-auth.js, player-identity.js", "Google sign-in, verified email matching, profile provisioning, authorization context."),
        ("Data access", "firebase-client.js, firebase-data.js", "Firebase initialization, public/operational reads, UI data hydration."),
        ("Lineup workflow", "lineup-submit.js, lineup-approve.js, lineup-reset.js, lineup-workflow-client.js", "Draft, submit, approve, reject, publish, and reset transactions."),
        ("Match operations", "match-management.js, score-rules.js", "Schedule entry, score validation, points, completion locking."),
        ("Season administration", "season-*.js, roster-admin-v3.js, venue-admin.js", "Season import, teams, matchups, rosters, venues, reset."),
        ("Public projection", "public-season-dashboard.js", "Publishes guest-safe season data after authorized operational changes."),
        ("Offline shell", "service-worker.js, manifest.webmanifest, pwa.js", "Installability, cache lifecycle, offline navigation fallback."),
        ("Security policy", "firestore.rules, firestore.indexes.json", "Authorization, transition constraints, query support."),
    ]
    add_table(doc, ["Component", "Primary modules", "Responsibility"], components, [1750, 3100, 4510], font_size=8.3)

    add_heading(doc, "4.2 Routing and module lifecycle", 2)
    add_body(
        doc,
        "The application is a hash-routed SPA. Views are declared as sections in index.html and activated by app.js. The runtime loader imports Firebase core modules and loads route-specific administrative modules only when their route or admin panel becomes active. Separate small bootstrap modules initialize lineup, approval, reset, EC status, and match-management features.",
    )
    add_body(
        doc,
        "Code assets use explicit query-string versions. Firebase Hosting sets no-cache headers for JavaScript and CSS. The service worker uses a versioned cache, applies network-first behavior to app code, and deletes earlier cache versions during activation.",
    )

    add_page_break(doc)
    add_heading(doc, "5. Deployment and runtime topology", 1)
    deployment = add_diagram(
        "Production deployment topology",
        [
            (40, 250, 350, 405, "User Device\nBrowser / Installed PWA", PALE_GREEN),
            (535, 95, 965, 245, "Firebase Hosting\nStatic files + SPA rewrite", PALE_BLUE),
            (535, 410, 965, 560, "Service Worker Cache\nApp shell + offline fallback", LIGHT),
            (1180, 95, 1540, 245, "Firebase Authentication\nGoogle provider", PALE_BLUE),
            (1180, 410, 1540, 560, "Cloud Firestore (nam5)\nRules + Indexes + Data", PALE_BLUE),
        ],
        [
            (350, 300, 535, 170, "HTTPS"),
            (350, 355, 535, 485, "Cache"),
            (965, 170, 1180, 170, "SDK"),
            (965, 485, 1180, 485, "SDK"),
        ],
        "deployment_topology.png",
    )
    add_picture_with_alt(
        doc,
        deployment,
        "Production deployment topology connecting the user device and service worker to Firebase Hosting, Authentication, and Firestore.",
        Inches(6.5),
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(caption.add_run("Figure 3. Current Spark-plan production topology"), size=9, italic=True, color=MUTED)

    topology_rows = [
        ("Firebase project", "alphaopen-development-2026", "Single project currently serves production."),
        ("Hosting URL", "alphaopen-development-2026.web.app", "Public endpoint; all paths rewrite to index.html."),
        ("Firestore", "(default), nam5", "Operational, private, and public documents."),
        ("Functions", "Not configured for deployment", "Functions source is retained for future migration but excluded from Spark deployment."),
        ("Headers", "No-cache for HTML, JS, CSS, manifest", "Reduces stale-code risk; service worker provides controlled cache fallback."),
        ("Emulators", "Auth 9099, Firestore 8080, Hosting 5000", "Local validation and security-rule testing."),
    ]
    add_table(doc, ["Element", "Configuration", "Architecture note"], topology_rows, [1850, 2800, 4710], font_size=8.8)

    add_heading(doc, "5.1 Environment and release management", 2)
    add_body(
        doc,
        "The repository currently targets one Firebase project and contains the Firebase client configuration directly in browser code, which is normal for Firebase clients. Production deployment is performed with the Firebase CLI for Hosting and Firestore Rules. There is no separate staging project, automated promotion workflow, or infrastructure-as-code pipeline.",
    )
    add_callout(
        doc,
        "Condition of approval",
        "Create separate development/test and production Firebase projects, automate rule tests in CI, and require reviewed deployment promotion before operational scale increases.",
        fill=PALE_GOLD,
        color=CAUTION,
    )

    add_page_break(doc)
    add_heading(doc, "6. Identity, authorization, and security architecture", 1)
    add_heading(doc, "6.1 Authentication and registration", 2)
    auth_numbering = begin_numbering(doc)
    for step in [
        "User signs in with Google through Firebase Authentication.",
        "The application requires a verified Google email.",
        "The normalized email must match an active Player Master email index.",
        "A new user record is created as pending and linked to the canonical Player ID.",
        "Super Admin approves the registration and assigns active season membership and roles.",
        "The application derives route access from global roles plus current-season member roles.",
    ]:
        add_number(doc, step, auth_numbering)

    add_heading(doc, "6.2 Authorization layers", 2)
    layers = [
        ("Navigation visibility", "data-access, data-access-any, data-nav-roles, data-super-admin", "User experience only; not a security boundary."),
        ("Client guard clauses", "Feature modules validate authorization context and team scope", "Prevents accidental actions and provides clear errors."),
        ("Firestore Rules", "Identity, active status, season role, team scope, allowed fields, and transition checks", "Primary security boundary."),
        ("Transaction invariants", "Current revision, current status, operation ID, actor UID, and multi-document state", "Concurrency and workflow integrity."),
        ("Immutable records", "Revision and review documents are create-only", "Nonrepudiation and audit preservation."),
    ]
    add_table(doc, ["Layer", "Mechanism", "Security role"], layers, [1800, 3900, 3660], font_size=8.5)

    add_heading(doc, "6.3 Role access matrix", 2)
    access_rows = [
        ("Public content", "Read", "Read", "Read", "Read", "Read", "Read/write"),
        ("Own player/account data", "-", "Read/update limited", "Read/update limited", "Read", "Read", "All"),
        ("Team lineup submit", "-", "-", "Own team", "Authorized scope", "Either team", "Any"),
        ("Lineup approve/reject", "-", "-", "-", "-", "Assigned scope", "Any"),
        ("Reset approved lineups", "-", "-", "-", "-", "Assigned scope", "Any"),
        ("Roster/season operations", "-", "-", "Request/view", "Season scope", "View", "Any"),
        ("Schedule/score", "Published read", "Assigned read", "Team scope", "Season scope", "Read", "Any"),
        ("Security administration", "-", "-", "-", "-", "-", "Any"),
    ]
    add_table(doc, ["Capability", "Guest", "Player", "Captain", "EC", "Approver", "Super Admin"], access_rows, [1900, 900, 1000, 1150, 1050, 1300, 2060], font_size=7.8)

    add_heading(doc, "6.4 Security controls and residual risk", 2)
    controls = [
        ("Verified identity", "Firebase ID token and verified Google email.", "Implemented"),
        ("Private data isolation", "Separate playerPrivate and venuePrivate documents.", "Implemented"),
        ("Role enforcement", "Rules read user, member, assignment, and matchup snapshots.", "Implemented"),
        ("Privilege snapshot", "Matchups carry captain and approver UID lists for efficient rule evaluation.", "Implemented"),
        ("Workflow actor binding", "Matchup stores operation ID and actor UID; related writes must match.", "Implemented"),
        ("Immutable workflow history", "Submitted revisions and action records cannot be updated or deleted.", "Implemented"),
        ("Exceptional reset score check", "Application queries line matches and transaction re-reads them.", "Application-enforced"),
        ("App Check", "No enforced App Check attestation.", "Gap"),
        ("Central monitoring", "No server event pipeline or alerting.", "Gap"),
        ("Bootstrap Super Admin", "Protected email is embedded in client/rules until custom claims are introduced.", "Accepted interim risk"),
    ]
    add_table(doc, ["Control", "Design", "Status"], controls, [2200, 5200, 1960], font_size=8.2)

    add_page_break(doc)
    add_heading(doc, "7. Data architecture", 1)
    add_heading(doc, "7.1 Collection hierarchy", 2)
    hierarchy = [
        ("Global identity", "users, players, playerPrivate, playerEmailIndex, playerAccountLinks, registrationRequests", "Account identity and canonical person records"),
        ("Global reference", "venues, venuePrivate, systemConfig, aoContent", "Venues, active season control, public content"),
        ("Season core", "seasons/{seasonId}", "Season configuration and status"),
        ("Season access", "members, approverAssignments", "Season roles and approval scope"),
        ("Competition setup", "teams, rosterSlots, rosterAssignments, ruleVersions, weeks, matchups", "Teams, ranks, rules, schedule structure"),
        ("Lineup workflow", "matchups/{matchupId}/lineups, revisions, lineupReviews", "Current lineups, immutable submissions, actions"),
        ("Match operations", "lineMatches, scheduleProposals, scoreSubmissions, scoreDecisions", "Line scheduling and scoring"),
        ("League operations", "availability, replacementRequests, latePassRequests, adjustments", "Supporting workflows"),
        ("Results", "standings, standingsSnapshots, playoffBrackets", "Calculated/public competition results"),
        ("Audit/notification", "auditEvents, users/{uid}/notifications", "Operational trace and user messages"),
    ]
    add_table(doc, ["Domain", "Collections", "Purpose"], hierarchy, [1800, 4300, 3260], font_size=8.1)

    add_heading(doc, "7.2 Identifier strategy", 2)
    identifiers = [
        ("Firebase UID", "Authentication account and role assignment"),
        ("Player ID", "Permanent person identity and historical linkage"),
        ("Season ID", "Partition key for competition operations"),
        ("Team ID", "Season team identity"),
        ("Assignment ID", "Immutable roster assignment interval"),
        ("Matchup ID", "Home-versus-Away team event"),
        ("Line match ID", "Matchup plus line number; normally {matchupId}-L1 through L5"),
        ("Operation ID", "Client-generated idempotency key for lineup action"),
        ("Revision number", "Monotonic lineup submission version within team/matchup"),
    ]
    add_table(doc, ["Identifier", "Use"], identifiers, [2300, 7060], font_size=9)

    add_heading(doc, "7.3 Public/private partitioning", 2)
    add_body(
        doc,
        "Public pages must read only explicitly published fields. Operational season documents include member UIDs, internal status, lineup workflow state, and administrative metadata. Personally identifiable information is split into private collections. Official records carry display-name snapshots for historical readability without requiring broad access to private master records.",
    )
    add_callout(
        doc,
        "Data-governance note",
        "Snapshot fields are intentionally denormalized. They preserve historical meaning but require defined correction procedures when an official record is wrong; changing the Player Master must not silently rewrite history.",
        fill=PALE_BLUE,
    )

    add_page_break(doc)
    add_heading(doc, "8. Core schema design", 1)
    add_heading(doc, "8.1 Identity and master data", 2)
    identity_schema = [
        ("users/{uid}", "uid, email, status, profileType, playerId, globalRoles, displayName, lastLoginAt", "Account authorization and profile"),
        ("players/{playerId}", "displayName, public fields, active status", "Canonical public player identity"),
        ("playerPrivate/{playerId}", "emailNormalized, phone, private notes, status", "Restricted PII and identity match"),
        ("playerAccountLinks/{playerId}", "uid, status, approval metadata, link method", "One-account-to-player link"),
        ("seasons/{seasonId}/members/{uid}", "roles, teamIds, playerId, status", "Season-scoped authorization"),
        ("teams/{teamId}", "name, captainPlayerIds, status", "Season team master"),
        ("rosterSlots/{team_rank}", "teamId, rankNumber, participation counters", "Persistent team-rank slot"),
        ("rosterAssignments/{assignmentId}", "teamId, rankNumber, playerId, type, dates, eligibility, status", "Effective-dated player assignment"),
    ]
    add_table(doc, ["Document", "Key fields", "Purpose"], identity_schema, [2500, 4650, 2210], font_size=7.9)

    add_heading(doc, "8.2 Matchup master", 2)
    matchup_schema = [
        ("Identity", "matchupId, weekId, stage, homeTeamId, awayTeamId"),
        ("Authorization records", "season member roles/teamIds and approver assignments"),
        ("Team states", "homeLineupStatus, awayLineupStatus"),
        ("Overall state", "lineupApprovalStatus, bothLineupsSubmitted, lineupsPublished"),
        ("Revision pointers", "homeLineupRevisionNumber, awayLineupRevisionNumber"),
        ("Actor tracking", "homeLineupTracking, awayLineupTracking"),
        ("Publication", "fullyApprovedAt, lineupsPublishedAt"),
        ("Concurrency/security", "lineupWorkflowActorUid, lineupWorkflowOperationId"),
        ("Reset", "approvalCycleNumber, lastLineupReset"),
        ("Competition", "status, completedLineCount, homeTeamPoints, awayTeamPoints, standingsApplied"),
    ]
    add_table(doc, ["Field group", "Fields"], matchup_schema, [2400, 6960], font_size=8.8)

    add_heading(doc, "8.3 Lineup current document and revision", 2)
    lineup_schema = [
        ("Identity", "seasonId, matchupId, teamId, revisionNumber, ruleVersionId"),
        ("State", "draft | submitted | approved | rejected"),
        ("Content", "five lines; each has two player IDs/names/ranks and SOR"),
        ("Validation", "passed, errors, checkedAt, checkedBy, ruleVersionId"),
        ("Submission", "submittedByUid/playerId/name/role, submittedAt"),
        ("Approval", "approvedByUid/playerId/name/role, approvedAt"),
        ("Rejection", "rejectionReason, rejectedByUid/playerId/name/role, rejectedAt"),
        ("Audit", "updatedByUid, updatedAt"),
        ("Revision copy", "same payload plus immutable=true in revisions/{revisionNumber}"),
    ]
    add_table(doc, ["Field group", "Definition"], lineup_schema, [2200, 7160], font_size=8.8)

    add_heading(doc, "8.4 Line match", 2)
    line_match_schema = [
        ("Identity", "lineMatchId, matchupId, lineNumber, seasonId"),
        ("Teams/players", "homeTeamId, awayTeamId, homePlayers[2], awayPlayers[2], rank snapshots"),
        ("Lineup linkage", "homeLineupRevisionNumber, awayLineupRevisionNumber, lineupState, scoreEntryAllowed"),
        ("Schedule", "venueId/name/address, scheduledAt, scheduleStatus"),
        ("Score", "sets, scoreStatus, winnerTeamId, homePoints, awayPoints, completedAt"),
        ("Audit", "createdAt, updatedAt"),
    ]
    add_table(doc, ["Field group", "Definition"], line_match_schema, [2200, 7160], font_size=8.8)

    add_page_break(doc)
    add_heading(doc, "9. Lineup submission and approval detailed design", 1)
    add_heading(doc, "9.1 Status model", 2)
    team_status = [
        ("pendingSubmission", "No sealed submission is available for this team.", "Captain/EC/Approver/Super Admin submits"),
        ("submitted", "Current revision is sealed and awaiting a decision.", "Approver approves or rejects"),
        ("approved", "Current team revision is approved.", "Other team approved or, before full approval, Approver rejects"),
        ("rejected", "Current revision was returned with a reason.", "Authorized submitter submits a new revision"),
    ]
    add_table(doc, ["Team status", "Meaning", "Next action"], team_status, [1900, 4300, 3160], font_size=8.8)
    matchup_status = [
        ("awaitingSubmission", "At least one team remains pending and neither team is rejected."),
        ("awaitingApproval", "Both teams are submitted/approved but at least one submitted revision needs approval."),
        ("rejected", "At least one team lineup is rejected."),
        ("fullyApproved", "Both team lineup statuses are approved."),
    ]
    add_table(doc, ["Matchup status", "Derivation"], matchup_status, [2200, 7160], font_size=9)

    add_heading(doc, "9.2 Submission transaction", 2)
    submission_numbering = begin_numbering(doc)
    for step in [
        "Validate request IDs, signed-in actor, role, team relationship, and five-line lineup.",
        "Read active roster assignments for the selected team and canonicalize player snapshots and ranks.",
        "Read the action, matchup, current lineup, and proposed revision inside a Firestore transaction.",
        "Reject replacement of a submitted or approved current lineup.",
        "Increment the revision number and write the current lineup plus immutable revision copy.",
        "Update the appropriate Home/Away status, revision pointer, tracking map, and derived matchup status.",
        "Set workflow actor UID and operation ID on the matchup.",
        "Create a write-once submitted action containing prior/new status and result.",
    ]:
        add_number(doc, step, submission_numbering)

    add_heading(doc, "9.3 Approval/rejection transaction", 2)
    approval_numbering = begin_numbering(doc)
    for step in [
        "Only a Neutral Approver assigned to the matchup/season or Super Admin may decide.",
        "The Approver may review one submitted team even while the opponent is pending.",
        "Approval requires the current status to be submitted; rejection permits submitted or approved before full approval.",
        "A rejection requires a reason and clears approval metadata.",
        "Self-approval is permitted for the current release and is recorded when submittedByUid equals approvedByUid.",
        "The matchup status is recalculated from both team states.",
        "When both teams become approved, five line-match records are created or refreshed and score entry is enabled.",
        "The decision and all publication writes are committed atomically with a create-only action record.",
    ]:
        add_number(doc, step, approval_numbering)

    add_heading(doc, "9.4 Visibility model", 2)
    visibility = [
        ("Captain", "Own team lineup", "Opponent team status only until publication"),
        ("EC", "Authorized season lineups", "Operational status and roster context"),
        ("Neutral Approver", "Both current team lineups", "Independent team decisions"),
        ("Super Admin", "All current and historical lineup data", "All actions"),
        ("Player/Guest", "Published information only", "No sealed drafts or internal decisions"),
    ]
    add_table(doc, ["Role", "Before full approval", "After full approval"], visibility, [1800, 3800, 3760], font_size=8.7)

    add_heading(doc, "9.5 Requirements variance from original specification", 2)
    variance = [
        ("Approver visibility", "Original: wait for both submissions", "Current: Approver can review the submitted team while the opponent is pending", "Approved product decision"),
        ("Approval granularity", "Original: approve both together", "Current: approve/reject each team independently; publish only when both approved", "Approved product decision"),
        ("Approver submission", "Original: not allowed", "Current: Neutral Approver may submit either team and self-approve with audit flag", "Approved product decision"),
        ("Time restrictions", "Original: deadlines envisioned", "Current: no approval time restrictions", "Explicitly deferred"),
        ("Post-approval change", "Original: change/consent workflow", "Current: no change request; separate reset-both exception", "Approved simplification"),
    ]
    add_table(doc, ["Topic", "Earlier baseline", "Current design", "Disposition"], variance, [1500, 2650, 3600, 1610], font_size=7.8)

    add_page_break(doc)
    add_heading(doc, "10. Reset Approved Lineup exception", 1)
    add_heading(doc, "10.1 Access and selection", 2)
    add_body(
        doc,
        "Reset Approved Lineup is a separate screen available only to Neutral Approvers and Super Admins. The user selects Season, Week, and Home Team/matchup. Both approved lineups, submitter/approver metadata, and all five line matches are displayed before reset.",
    )
    add_heading(doc, "10.2 Eligibility validation", 2)
    for item in [
        "Matchup must be fullyApproved, or both stored team statuses must be approved.",
        "Both current lineup documents must exist.",
        "The application queries every line match and blocks when it finds in-progress/completed states, submitted/confirmed/disputed/published score states, nonzero set scores, nonzero points, a winner, or completedAt.",
        "The reset transaction re-reads the discovered line documents plus the five canonical line IDs immediately before writing.",
        "The user must enter a reason, acknowledge the effect, and type the exact Matchup ID.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "10.3 Reset transaction result", 2)
    reset_rows = [
        ("Home lineup", "Delete current document; preserve revisions"),
        ("Away lineup", "Delete current document; preserve revisions"),
        ("Team statuses", "Both pendingSubmission"),
        ("Matchup status", "awaitingSubmission"),
        ("Publication", "lineupsPublished=false; timestamps cleared"),
        ("Approval cycle", "Increment by one"),
        ("Line matches", "Retain schedule documents; set lineupState=awaitingReapproval and scoreEntryAllowed=false"),
        ("Audit", "Create resetBothApprovedLineups action with actor, reason, previous revisions, and cycle"),
    ]
    add_table(doc, ["Record", "Result"], reset_rows, [2200, 7160], font_size=8.8)
    add_callout(
        doc,
        "Residual risk",
        "Firestore Rules cannot issue an authorization-time collection query proving that no scored line exists. The application performs the score inspection and transaction re-read. This is reasonable for a rare action restricted to trusted Approvers and Super Admins, but it is the strongest trigger for future trusted-server migration.",
        fill=PALE_RED,
        color=RISK_RED,
    )

    add_page_break(doc)
    add_heading(doc, "11. Scheduling, scoring, and publication", 1)
    add_heading(doc, "11.1 Match schedule and score", 2)
    add_body(
        doc,
        "Each fully approved matchup produces five operational line-match documents. Captains may update lines involving their teams; EC and Super Admin users have season-wide management scope. A reset line with scoreEntryAllowed=false or lineupState=awaitingReapproval cannot be scored until both new lineups are fully approved.",
    )
    add_heading(doc, "11.2 Score validation", 2)
    for item in [
        "Set scores are structured numeric values rather than free-form result text.",
        "The scoring module validates completeness and determines the winner.",
        "Straight-set completion does not permit a third set.",
        "Split first and second sets require a deciding set.",
        "Winner receives 14 league points; loser points follow the configured two-set or three-set rules.",
        "Final score entry requires explicit user confirmation and then disables further normal score update.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11.3 Public projection", 2)
    add_body(
        doc,
        "Operational writes are followed by publication of guest-safe season dashboard data. The public projection must contain only approved schedules, published lineups, results, standings, and approved venue fields. Private contact details, registration state, availability, sealed lineups, internal notes, and dispute information remain outside the public projection.",
    )

    add_heading(doc, "11.4 Consistency considerations", 2)
    consistency = [
        ("Within a workflow", "Firestore transaction/batch", "Atomic current state, revision, matchup, line records, and action."),
        ("Public projection", "Follow-on client publication", "Eventual consistency; a failed projection may leave public data stale."),
        ("Standings", "Client/server-derived documents depending workflow", "Must be idempotent and protected from duplicate application."),
        ("Offline writes", "Firestore client behavior", "Operational flows should surface connectivity and avoid assuming immediate publication."),
    ]
    add_table(doc, ["Boundary", "Mechanism", "Implication"], consistency, [1900, 2700, 4760], font_size=8.6)

    add_page_break(doc)
    add_heading(doc, "12. Transaction, concurrency, and integrity design", 1)
    integrity = [
        ("Idempotency", "Every official lineup operation carries a random operation ID; duplicate action reads return the prior result."),
        ("Optimistic concurrency", "Firestore transaction retries when read documents change."),
        ("Revision monotonicity", "Next revision is max(current lineup revision, matchup pointer) + 1."),
        ("Sealed state", "Submitted and approved current lineups cannot be overwritten by normal submission."),
        ("Cross-document binding", "Matchup actor UID and operation ID bind related write permissions in Rules."),
        ("State derivation", "Matchup approval status is recalculated from Home and Away statuses."),
        ("Immutable evidence", "Revisions and review actions are create-only."),
        ("Reset safety", "Known line documents are re-read before reset; score activity causes application failure."),
    ]
    add_table(doc, ["Integrity concern", "Control"], integrity, [2300, 7060], font_size=8.8)

    add_heading(doc, "12.1 Rule-evaluation budget", 2)
    add_body(
        doc,
        "Firestore Rules have expression and document-access limits. Workflow rules use matchup authorization snapshots and a single operation actor binding to avoid repeatedly loading user, member, and assignment records for every write in a multi-document transaction. Emulator tests validate the complete submission, one-sided approval, second-team approval with five line records, unauthorized Captain denial, and reset transaction.",
    )

    add_heading(doc, "12.2 Failure behavior", 2)
    failure_rows = [
        ("Validation failure", "No official write; display exact lineup or score error."),
        ("Authorization failure", "Firestore permission-denied; UI restores controls and displays failure."),
        ("Concurrent update", "Firestore retries transaction; stale state is rejected on re-evaluation."),
        ("Partial network failure", "Atomic transaction does not partially commit."),
        ("Public projection failure", "Operational state may commit while public projection remains stale; requires retry/monitoring."),
        ("Service worker stale asset", "Network-first app code and no-cache hosting headers reduce risk; versioned cache activates on update."),
    ]
    add_table(doc, ["Failure", "Expected behavior"], failure_rows, [2300, 7060], font_size=8.8)

    add_page_break(doc)
    add_heading(doc, "13. Nonfunctional architecture", 1)
    nfr_rows = [
        ("Availability", "Firebase managed services; static hosting remains available independently of custom servers.", "Define service objective and status monitoring."),
        ("Performance", "Static CDN delivery, lazy modules, Firestore direct reads, indexed queries.", "Add query budgets and representative device testing."),
        ("Scalability", "Adequate for a small league; denormalized documents and direct reads scale horizontally.", "Review read amplification and public projection strategy before multi-league expansion."),
        ("Security", "Firebase Auth + Rules + immutable records.", "Add App Check, environment separation, custom claims/server controls."),
        ("Privacy", "Public/private collections and limited contact access.", "Document retention, deletion, and incident procedures."),
        ("Accessibility", "Responsive semantic HTML and mobile layout.", "Perform formal WCAG 2.2 AA audit and keyboard/screen-reader testing."),
        ("Maintainability", "Feature modules and shared workflow service.", "Introduce build tooling, linting, typed contracts, and module tests."),
        ("Observability", "Browser errors and Firebase console.", "Add structured event/error monitoring and deployment traceability."),
        ("Recoverability", "Firestore managed durability; immutable history.", "Schedule exports/backups and test restore."),
        ("Compatibility", "Modern iOS/Android/desktop browsers; PWA manifest.", "Maintain supported-browser matrix."),
    ]
    add_table(doc, ["Quality attribute", "Current design", "Required control"], nfr_rows, [1600, 3900, 3860], font_size=8.1)

    add_heading(doc, "13.1 Data retention and recovery", 2)
    for item in [
        "Submitted lineup revisions and official workflow actions should be retained for the full league-history period.",
        "Private player data should follow a documented retention and correction policy.",
        "Season reset must preserve the season root and immutable audit evidence while deleting explicitly scoped descendants.",
        "Production Firestore export and restore procedures should be exercised before the next season launch.",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "14. Testing and architecture assurance", 1)
    assurance = [
        ("Static application contract", "test-prototype.mjs", "Routes, views, modules, PWA, encoding, score rules, lineup service usage."),
        ("Workflow logic", "functions/test/workflow.test.js", "Status derivation, five-line validation, score-activity detection."),
        ("Rules compilation", "Firebase Firestore emulator", "Production rules compile without errors."),
        ("Role/workflow integration", "spark-workflow-rules-test.mjs", "Submission, Captain approval denial, one-sided approval, full approval, reset."),
        ("Browser smoke test", "Local hosted application", "Required views and versioned modules load without console errors."),
        ("Production verification", "Firebase deploy output and HTTPS request", "Hosting/rules released; live response is HTTP 200."),
    ]
    add_table(doc, ["Assurance layer", "Mechanism", "Coverage"], assurance, [1900, 3000, 4460], font_size=8.6)

    add_heading(doc, "14.1 Required pre-release gates", 2)
    for item in [
        "All JavaScript syntax and static contract tests pass.",
        "Firestore Rules compile and the Spark workflow emulator test passes.",
        "Rules changes receive security-focused peer review.",
        "Representative Captain, Approver, EC, and Super Admin user journeys pass in a non-production Firebase project.",
        "PWA cache version and module query versions are advanced when deployable assets change.",
        "Production smoke test confirms route load, authentication, and read/write permission behavior.",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "15. Risks, conditions, and technical debt", 1)
    risks = [
        ("R1", "High", "Privileged workflow logic executes in browser", "Direct SDK callers may attempt crafted writes; Rules are the only enforcement tier.", "Retain emulator tests; migrate protected mutations to Cloud Functions when trigger is met."),
        ("R2", "High", "Exceptional reset score check is application-enforced", "A deliberately modified privileged client could omit a scored line.", "Trusted-role restriction now; move validation to server on Blaze."),
        ("R3", "High", "Single Firebase environment", "Testing or administrative mistakes can affect production data.", "Create separate dev/test and production projects."),
        ("R4", "Medium", "Bootstrap Super Admin email embedded", "Long-lived special-case privilege and account concentration.", "Provision custom claim or server-managed global role; maintain two-person admin coverage."),
        ("R5", "Medium", "Manual deployment and cache versioning", "Human error may release mismatched assets or rules.", "CI/CD, manifest generation, deployment approval, rollback procedure."),
        ("R6", "Medium", "Public projection is follow-on client work", "Operational and public views may temporarily diverge.", "Add retry dashboard now; move projection to event-driven backend later."),
        ("R7", "Medium", "No App Check or centralized telemetry", "Abuse and client failures are harder to detect.", "Enable App Check and error/event monitoring."),
        ("R8", "Medium", "Specification drift", "Earlier specification conflicts with current team-level approval decisions.", "Approve this document as the architecture baseline and update product specification."),
        ("R9", "Low", "CDN-loaded SDK and unbundled modules", "Dependency or browser-loading behavior is less controlled.", "Introduce lockfile/build pipeline and Subresource Integrity where feasible."),
        ("R10", "Low", "No formal WCAG audit", "Usability barriers may remain.", "Perform WCAG 2.2 AA review before broad launch."),
    ]
    add_table(doc, ["ID", "Rating", "Risk", "Impact", "Treatment"], risks, [600, 800, 2350, 2450, 3160], font_size=7.3)

    add_heading(doc, "15.1 Conditions of approval", 2)
    conditions = [
        ("C1", "Approve and publish the updated product/design baseline reflecting team-level approval and reset-both behavior.", "Product owner", "Before next season configuration"),
        ("C2", "Create a non-production Firebase project and run role-based end-to-end tests there.", "Engineering", "Before next material workflow change"),
        ("C3", "Add CI gates for static tests and Firestore Rules emulator tests.", "Engineering", "Before multi-contributor development"),
        ("C4", "Document backup/export and restore procedures.", "Operations", "Before production season launch"),
        ("C5", "Define the Blaze migration trigger and budget owner.", "ARB/Product", "At architecture approval"),
        ("C6", "Review bootstrap admin and App Check strategy.", "Security owner", "Within one release cycle"),
    ]
    add_table(doc, ["ID", "Condition", "Owner", "Target"], conditions, [650, 5200, 1500, 2010], font_size=7.9)

    add_page_break(doc)
    add_heading(doc, "16. Future Blaze migration architecture", 1)
    add_body(
        doc,
        "The application already isolates official lineup mutations behind lineup-workflow-client.js. The Firestore schema, UI payloads, status values, operation IDs, revisions, and audit documents can remain unchanged. A Blaze upgrade replaces the transaction implementation with callable Cloud Functions or HTTPS endpoints using the Firebase Admin SDK.",
    )
    migration = add_diagram(
        "Controlled migration from Spark to trusted services",
        [
            (45, 130, 405, 275, "Current UI\nUnchanged", PALE_GREEN),
            (45, 420, 405, 565, "Current Firestore Schema\nUnchanged", PALE_BLUE),
            (620, 130, 1030, 275, "Workflow Service Interface\nSame operation payloads", LIGHT),
            (620, 420, 1030, 565, "Tightened Rules\nClient reads; server writes", PALE_GOLD),
            (1240, 130, 1555, 275, "Callable Functions\nTrusted validation", PALE_BLUE),
            (1240, 420, 1555, 565, "Admin SDK\nAtomic transactions", PALE_BLUE),
        ],
        [
            (405, 200, 620, 200, ""),
            (1030, 200, 1240, 200, ""),
            (1400, 275, 1400, 420, ""),
            (1240, 490, 1030, 490, ""),
            (620, 490, 405, 490, ""),
        ],
        "blaze_migration.png",
    )
    add_picture_with_alt(
        doc,
        migration,
        "Future migration path from the unchanged user interface through a workflow service interface to callable functions, Admin SDK transactions, and tightened Firestore Rules.",
        Inches(6.5),
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(caption.add_run("Figure 4. Future trusted mutation path"), size=9, italic=True, color=MUTED)

    add_heading(doc, "16.1 Migration steps", 2)
    migration_steps = [
        "Enable Blaze billing and required Cloud Functions, Cloud Build, and Artifact Registry APIs.",
        "Deploy callable operations for list seasons, submit lineup, decide team lineup, and reset both lineups.",
        "Move actor resolution, roster canonicalization, transition validation, score inspection, and publication into trusted code.",
        "Switch the workflow service implementation from Firestore transactions to callable requests.",
        "Tighten Rules so official workflow documents are read-only to browser clients; retain draft access where appropriate.",
        "Enable Firebase App Check enforcement after monitoring legitimate traffic.",
        "Add event-driven public projection, notification, and audit functions as separate idempotent handlers.",
        "Run dual-path test fixtures and cut over without data migration.",
    ]
    migration_numbering = begin_numbering(doc)
    for step in migration_steps:
        add_number(doc, step, migration_numbering)

    add_heading(doc, "16.2 Recommended migration triggers", 2)
    triggers = [
        ("Security", "ARB requires server-side proof of no score activity or stronger nonrepudiation."),
        ("Scale", "Multiple leagues/seasons, materially more users, or increasing direct-client abuse risk."),
        ("Automation", "Need scheduled deadlines, notifications, penalties, projections, or background reconciliation."),
        ("Operations", "Need centralized audit events, alerting, retry queues, or service-level objectives."),
        ("Integration", "Need email/SMS, payments, court systems, or third-party APIs requiring secrets."),
    ]
    add_table(doc, ["Trigger class", "Threshold"], triggers, [2000, 7360], font_size=8.8)

    add_page_break(doc)
    add_heading(doc, "17. Open architecture decisions for ARB", 1)
    open_items = [
        ("A1", "Is the Spark-safe privileged-client model acceptable through the next season?", "Approve interim / require Blaze before launch"),
        ("A2", "Is self-approval by a Neutral Approver acceptable when explicitly recorded?", "Accept / require separate approver"),
        ("A3", "Should a Neutral Approver be season-wide or restricted to matchup assignments?", "Season / week / matchup policy"),
        ("A4", "What event triggers mandatory Blaze migration?", "Date, user volume, security control, or feature dependency"),
        ("A5", "What are retention periods for private player data and immutable competition records?", "Retention schedule"),
        ("A6", "What RTO/RPO applies to season operations?", "Recovery objectives"),
        ("A7", "Who owns production deployment approval and emergency rollback?", "Named operational owner"),
        ("A8", "Which earlier specification clauses are formally superseded by this document?", "Baseline reconciliation"),
    ]
    add_table(doc, ["ID", "Question", "Required decision"], open_items, [700, 5900, 2760], font_size=8.2)

    add_heading(doc, "17.1 Proposed ARB disposition", 2)
    add_callout(
        doc,
        "Proposed decision",
        "Approve with conditions C1-C6. Permit Spark-safe operation for the current trusted league-administration scope. Require a Blaze migration before introducing untrusted delegated administrators, secret-bearing integrations, automated deadline enforcement, or a security requirement for server-side reset validation.",
        fill=PALE_GREEN,
        color=GREEN,
    )

    add_page_break(doc)
    add_heading(doc, "Appendix A. Lineup state transition reference", 1)
    transitions = [
        ("pendingSubmission", "submitted", "Captain, EC, Approver, Super Admin", "Valid five-line roster; new revision and action"),
        ("rejected", "submitted", "Captain, EC, Approver, Super Admin", "Corrected valid lineup; new revision"),
        ("submitted", "approved", "Approver, Super Admin", "Current revision; decision audit"),
        ("submitted", "rejected", "Approver, Super Admin", "Required reason"),
        ("approved", "rejected", "Approver, Super Admin", "Only before matchup is fullyApproved"),
        ("approved + approved", "fullyApproved", "System transaction", "Publish both and enable five line matches"),
        ("fullyApproved", "awaitingSubmission", "Approver, Super Admin", "Reset both, no detected score activity, reason and exact confirmation"),
    ]
    add_table(doc, ["From", "To", "Actor", "Guard/effect"], transitions, [1700, 1750, 2350, 3560], font_size=8.2)

    add_heading(doc, "Appendix B. Firestore Rules ownership summary", 1)
    rule_ownership = [
        ("users", "Owner limited profile updates; Super Admin account management"),
        ("players / playerPrivate", "Public active player read vs restricted PII"),
        ("seasons / members", "Active member reads; Super Admin writes"),
        ("teams / rosters / weeks", "Season members read; EC/Super Admin write"),
        ("matchups", "Workflow fields through validated role/status transitions; admin structural writes"),
        ("lineups", "Team-scoped draft/submission; reviewer decisions through operation binding"),
        ("revisions / lineupReviews", "Create-only official evidence"),
        ("lineMatches", "Reviewer publication/reset fields; Captain/EC schedule and score fields"),
        ("approverAssignments", "Authorized self/EC reads; Super Admin writes"),
        ("standings / snapshots", "Member reads; protected calculated writes"),
        ("auditEvents", "EC reads; protected writes"),
        ("unmatched paths", "Deny by default"),
    ]
    add_table(doc, ["Path", "Ownership rule"], rule_ownership, [3000, 6360], font_size=8.7)

    add_heading(doc, "Appendix C. Architecture evidence", 1)
    evidence = [
        ("ALPHAOPEN_APP_SPEC.md", "Product scope, roles, nonfunctional requirements, original workflow baseline"),
        ("FIRESTORE_DATA_MODEL.md", "Canonical schema, identifiers, collections, ownership, indexes"),
        ("firebase.json", "Hosting, Firestore, emulator, and deployment configuration"),
        ("firestore.rules", "Authorization and workflow transition enforcement"),
        ("firestore.indexes.json", "Composite query support"),
        ("firebase-auth.js", "Authentication, registration, profile, and role derivation"),
        ("runtime-loader.js / service-worker.js", "Module lifecycle, caching, and PWA runtime"),
        ("lineup-workflow-client.js", "Spark-safe official workflow transactions"),
        ("lineup-submit.js / lineup-approve.js / lineup-reset.js", "Role-specific UI workflows"),
        ("match-management.js / score-rules.js", "Scheduling, scoring, locking, and points"),
        ("test-prototype.mjs / emulator tests", "Architecture and security assurance evidence"),
    ]
    add_table(doc, ["Artifact", "Evidence"], evidence, [3100, 6260], font_size=8.5)

    add_heading(doc, "Appendix D. Glossary", 1)
    glossary = [
        ("ARB", "Architecture Review Board"),
        ("EC", "Executive Committee"),
        ("PWA", "Progressive Web Application"),
        ("SOR", "Sum of roster ranks for the two players on a doubles line"),
        ("Spark", "Firebase no-cost plan used by the current deployment"),
        ("Blaze", "Firebase pay-as-you-go plan required for the planned Cloud Functions tier"),
        ("Current lineup", "Mutable pointer/document representing the latest team revision and state"),
        ("Revision", "Immutable copy of a submitted lineup"),
        ("Line match", "One of five doubles matches within a team matchup"),
        ("Public projection", "Guest-safe derived document set published from operational data"),
    ]
    add_table(doc, ["Term", "Definition"], glossary, [2200, 7160], font_size=8.8)

    # Core properties
    props = doc.core_properties
    props.title = "AlphaOpen Architecture Review Board Detailed Design"
    props.subject = "Application design, schemas, security, workflows, deployment, risks, and future-state architecture"
    props.author = "AlphaOpen Architecture"
    props.keywords = "AlphaOpen, Firebase, Firestore, PWA, Architecture Review Board, lineup approval"
    props.comments = "Generated from the deployed application architecture and repository design artifacts."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
